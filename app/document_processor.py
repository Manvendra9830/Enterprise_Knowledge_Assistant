"""
Document Processor — Handles loading, text extraction, chunking, and metadata management.

Supports PDF, DOCX, TXT, and Markdown files. Produces DocumentChunk objects
with page-level metadata for downstream embedding and indexing.

Design decisions:
- Chunk size ~500 chars with 50 char overlap balances retrieval precision with context.
- Paragraph-boundary-aware splitting avoids breaking sentences mid-thought.
- Page tracking enables precise source citation in answers.
"""

from __future__ import annotations

import hashlib
import logging
import os
import re
from pathlib import Path
from typing import Optional

from app.models import DocumentChunk

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Loads, extracts text, and chunks documents with metadata."""

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def process_file(self, file_path: str, file_content: Optional[bytes] = None) -> list[DocumentChunk]:
        """Process a single file and return a list of DocumentChunks."""
        path = Path(file_path)
        file_type = path.suffix.lower().lstrip(".")
        document_name = path.name

        logger.info(f"Processing document: {document_name} (type: {file_type})")

        # Extract text with page information
        pages = self._extract_text(file_path, file_type, file_content)

        if not pages:
            logger.warning(f"No text extracted from {document_name}")
            return []

        # Chunk each page and collect all chunks
        chunks: list[DocumentChunk] = []
        global_chunk_index = 0

        for page_num, page_text in pages:
            page_chunks = self._chunk_text(page_text)
            for chunk_text in page_chunks:
                chunk_id = self._generate_chunk_id(document_name, page_num, global_chunk_index)
                chunks.append(
                    DocumentChunk(
                        chunk_id=chunk_id,
                        text=chunk_text.strip(),
                        document_name=document_name,
                        page_number=page_num,
                        chunk_index=global_chunk_index,
                        file_type=file_type,
                        metadata={
                            "source": document_name,
                            "page": page_num,
                            "chunk_index": global_chunk_index,
                        },
                    )
                )
                global_chunk_index += 1

        logger.info(f"Produced {len(chunks)} chunks from {document_name}")
        return chunks

    def process_directory(self, directory: str) -> list[DocumentChunk]:
        """Process all supported files in a directory."""
        supported = {".pdf", ".docx", ".txt", ".md"}
        all_chunks: list[DocumentChunk] = []

        dir_path = Path(directory)
        if not dir_path.exists():
            logger.error(f"Directory not found: {directory}")
            return all_chunks

        for file_path in sorted(dir_path.iterdir()):
            if file_path.suffix.lower() in supported and file_path.is_file():
                try:
                    chunks = self.process_file(str(file_path))
                    all_chunks.extend(chunks)
                except Exception as e:
                    logger.error(f"Error processing {file_path.name}: {e}")

        logger.info(f"Total chunks from directory: {len(all_chunks)}")
        return all_chunks

    # ------------------------------------------------------------------
    # Text Extraction (per format)
    # ------------------------------------------------------------------

    def _extract_text(
        self, file_path: str, file_type: str, file_content: Optional[bytes] = None
    ) -> list[tuple[int, str]]:
        """Return list of (page_number, text) tuples."""
        extractors = {
            "pdf": self._extract_pdf,
            "docx": self._extract_docx,
            "txt": self._extract_txt,
            "md": self._extract_txt,  # Markdown treated as plain text
        }
        extractor = extractors.get(file_type)
        if extractor is None:
            logger.warning(f"Unsupported file type: {file_type}")
            return []
        return extractor(file_path, file_content)

    def _extract_pdf(self, file_path: str, file_content: Optional[bytes] = None) -> list[tuple[int, str]]:
        """Extract text from PDF with page-level granularity."""
        try:
            from pdfminer.high_level import extract_pages
            from pdfminer.layout import LTTextContainer

            pages = []
            source = file_path

            # If file_content is provided, write to temp and read
            if file_content:
                import tempfile
                tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
                tmp.write(file_content)
                tmp.close()
                source = tmp.name

            for page_num, page_layout in enumerate(extract_pages(source), start=1):
                page_text = ""
                for element in page_layout:
                    if isinstance(element, LTTextContainer):
                        page_text += element.get_text()
                if page_text.strip():
                    pages.append((page_num, page_text))

            # Clean up temp file
            if file_content:
                os.unlink(source)

            return pages
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            return []

    def _extract_docx(self, file_path: str, file_content: Optional[bytes] = None) -> list[tuple[int, str]]:
        """Extract text from DOCX. DOCX doesn't have native pages, so we treat the whole doc as page 1."""
        try:
            from docx import Document
            import io

            if file_content:
                doc = Document(io.BytesIO(file_content))
            else:
                doc = Document(file_path)

            full_text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
            if full_text.strip():
                return [(1, full_text)]
            return []
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return []

    def _extract_txt(self, file_path: str, file_content: Optional[bytes] = None) -> list[tuple[int, str]]:
        """Extract text from TXT/MD files. Simulate pages using section headers."""
        try:
            if file_content:
                text = file_content.decode("utf-8", errors="replace")
            else:
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    text = f.read()

            if not text.strip():
                return []

            # Split into sections based on top-level markdown headers (## )
            sections = re.split(r"\n(?=## )", text)
            pages = []
            for i, section in enumerate(sections, start=1):
                if section.strip():
                    pages.append((i, section))

            # If no sections found, return as single page
            if not pages:
                pages = [(1, text)]

            return pages
        except Exception as e:
            logger.error(f"Text extraction error: {e}")
            return []

    # ------------------------------------------------------------------
    # Chunking
    # ------------------------------------------------------------------

    def _chunk_text(self, text: str) -> list[str]:
        """
        Split text into chunks using paragraph-aware recursive splitting.

        Strategy:
        1. Split by double newlines (paragraphs) first
        2. If a paragraph exceeds chunk_size, split by single newlines
        3. If still too large, split by sentences
        4. Apply overlap between chunks
        """
        if not text or not text.strip():
            return []

        # Split into paragraphs
        paragraphs = re.split(r"\n\s*\n", text)
        paragraphs = [p.strip() for p in paragraphs if p.strip()]

        chunks: list[str] = []
        current_chunk = ""

        for para in paragraphs:
            # If adding this paragraph would exceed chunk size
            if len(current_chunk) + len(para) + 1 > self.chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                # Overlap: keep the tail of current chunk
                overlap_text = current_chunk[-self.chunk_overlap:] if self.chunk_overlap > 0 else ""
                current_chunk = overlap_text + " " + para if overlap_text else para
            else:
                current_chunk = current_chunk + "\n\n" + para if current_chunk else para

            # If a single paragraph is too large, force-split it
            if len(current_chunk) > self.chunk_size * 1.5:
                sub_chunks = self._force_split(current_chunk)
                for sc in sub_chunks[:-1]:
                    chunks.append(sc.strip())
                current_chunk = sub_chunks[-1] if sub_chunks else ""

        if current_chunk.strip():
            chunks.append(current_chunk.strip())

        return chunks

    def _force_split(self, text: str) -> list[str]:
        """Force-split a long text by sentences, then by character limit."""
        sentences = re.split(r"(?<=[.!?])\s+", text)
        chunks: list[str] = []
        current = ""

        for sentence in sentences:
            if len(current) + len(sentence) + 1 > self.chunk_size and current:
                chunks.append(current.strip())
                overlap = current[-self.chunk_overlap:] if self.chunk_overlap > 0 else ""
                current = overlap + " " + sentence if overlap else sentence
            else:
                current = current + " " + sentence if current else sentence

        if current.strip():
            chunks.append(current.strip())

        return chunks

    # ------------------------------------------------------------------
    # Utilities
    # ------------------------------------------------------------------

    @staticmethod
    def _generate_chunk_id(doc_name: str, page: int, chunk_idx: int) -> str:
        """Generate a deterministic chunk ID."""
        raw = f"{doc_name}::page{page}::chunk{chunk_idx}"
        return hashlib.sha256(raw.encode()).hexdigest()[:16]
